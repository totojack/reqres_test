#!/usr/bin/env python
# Requires Python 3.6 or above.

import requests
import os
import json
from docx import Document
from docx.text.paragraph import Paragraph

def get_data():
    """It manages the data retrieval."""

    print("Retrieving data...")
    url = os.environ['REST_API']

    #first page processing
    response = requests.get(url, timeout=(5, 10))
    json_res = json.loads(response.text)
    #print(json.dumps(json_res, indent = 4, sort_keys=True))

    #first data
    result = json_res["data"]

    #iterate the other pages
    total_pages = json_res["total_pages"]
    print(f"There are {total_pages} pages")

    if (total_pages > 1):

        for i in range(2,total_pages+1):
            url =  os.environ['REST_API'] + "?page=" + str(i)
            print(f"Retrieving url '{url}'...")
            response = requests.get(url, timeout=(5, 10))
            json_res_pg = json.loads(response.text)
            #print(json.dumps(json_res, indent = 4, sort_keys=True))

            #add new data to final object
            result += json_res_pg["data"]

    #print(result)
    return result    

def generate_doc(my_data):
    """It generates the doc"""

    document = Document()
    document.add_heading('Users list', 0)
    p = document.add_paragraph()

    for item in my_data:
        #print(item)
        p.add_run(item['first_name'] + " " + item['last_name']).add_break()

    document.save('/out/output.docx')

def main():
    """get data"""
    my_data = get_data()

    """generate doc"""
    generate_doc(my_data)

if __name__ == "__main__":
    main()
